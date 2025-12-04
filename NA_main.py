# 原作者：杨天乐@关西大学 / Author: Shiame Yeung@Kansai University / 作成者：楊　天楽@関西大学
#!/usr/bin/env python3
# coding: utf-8
"""
na_pipeline.py  ——  单文件版（Step‑1 对齐 + 扩展公司识别）
2025‑07‑08  rev‑C
"""

def cute_box(cn: str, jp: str, icon: str = "🌸") -> None:
    """
    多行也能对齐的可爱中/日双语框
    cn: 中文提示（可以多行，用 '\\n' 分隔）
    jp: 日文提示（可以多行）
    icon: 每行开头和结尾的小表情
    """
    # 把中/日各自的多行拆开，拼成统一列表
    lines = []
    for segment in (cn, jp):
        for ln in segment.split("\n"):
            ln = ln.strip()
            # 用 "icon + 空格 + 文本 + 空格 + icon" 构造每一行
            lines.append(f"{icon} {ln} {icon}")

    # 找到最长那行，做为框宽
    width = max(len(ln) for ln in lines)
    border = "─" * width

    # 打印上边框
    print(f"╭{border}╮")
    # 打印每一行，右侧填充空格到 width
    for ln in lines:
        print("│" + ln.ljust(width) + "│")
    # 打印下边框
    print(f"╰{border}╯")

import sys, subprocess, os

def ensure_env() -> None:
    """
    环境自检与自动修复程序
    1. 检查所有必要的库 (包括 OpenAI, GLiNER, RapidFuzz 等)
    2. 缺失则自动调用 pip 安装
    3. 安装完成后自动重启脚本，实现无缝体验
    """
    import sys
    import subprocess
    import pkg_resources
    from pkg_resources import DistributionNotFound, VersionConflict

    # --- 定义项目所需的全部依赖 ---
    # 格式遵循 pip requirements.txt 标准
    REQUIRED_PACKAGES = [
        # 基础工具
        "pandas", 
        "tqdm", 
        "requests",
        "packaging",
        
        # 数据库
        "sqlalchemy", 
        "pymysql",
        
        # 文本处理
        "python-docx", 
        "rapidfuzz",  # 模糊匹配
        
        # AI 与 模型 (核心)
        "openai>=1.0.0",          # 必须 1.0 以上版本
        "gliner",                 # 新增：实体提取
        "sentence-transformers",  # 语义向量
        "torch",                  # 深度学习后端
        "transformers",           # HuggingFace 工具
        
        # 旧版兼容 (如果还用 spaCy)
        "spacy",
    ]

    # 检查当前 Python 版本以决定特定依赖 (可选)
    py_major, py_minor = sys.version_info[:2]
    if (py_major, py_minor) >= (3, 13):
        # Python 3.13+ 可能需要特定版本的 numpy 或其他库，这里暂且保留通用
        pass

    missing = []
    
    # --- 1. 检查缺失包 ---
    for pkg in REQUIRED_PACKAGES:
        try:
            pkg_resources.require(pkg)
        except (DistributionNotFound, VersionConflict):
            missing.append(pkg)

    # --- 2. 检查 spaCy 模型 (特例) ---
    try:
        import spacy
        if not spacy.util.is_package("en_core_web_sm"):
            missing.append("spacy_model:en_core_web_sm")
    except ImportError:
        pass # spacy 本身缺失会在上面被捕获

    # --- 3. 执行安装 ---
    if missing:
        cute_box(
            f"检测到缺失依赖，正在自动安装...\n缺失项: {', '.join(missing)}",
            f"不足している依存関係を検出しました。自動インストール中...\n対象: {', '.join(missing)}",
            "📦"
        )
        
        # 分离普通包和 spaCy 模型
        pip_pkgs = [p for p in missing if not p.startswith("spacy_model:")]
        spacy_models = [p for p in missing if p.startswith("spacy_model:")]

        # 安装 pip 包
        if pip_pkgs:
            try:
                # 注意：这里去掉了 stdout=subprocess.DEVNULL，让用户看到进度条
                subprocess.check_call([sys.executable, "-m", "pip", "install"] + pip_pkgs)
            except subprocess.CalledProcessError as e:
                cute_box(f"安装失败: {e}\n请尝试手动运行: pip install {' '.join(pip_pkgs)}", 
                         "インストールに失敗しました。手動で実行してください。", "❌")
                sys.exit(1)

        # 安装 spaCy 模型
        for model in spacy_models:
            model_name = model.split(":")[1]
            print(f"⬇️ Downloading spaCy model: {model_name}...")
            subprocess.check_call([sys.executable, "-m", "spacy", "download", model_name])

        cute_box(
            "依赖安装完成！正在自动重启程序...",
            "インストール完了！プログラムを自動再起動します...",
            "🔄"
        )

        # --- 4. 自动重启脚本 (黑科技) ---
        # 使用 os.execv 重新加载当前脚本，继承当前的进程 ID
        # 这样用户就不需要手动再输一次命令了
        os.execv(sys.executable, [sys.executable] + sys.argv)

# —————— 在脚本一启动就先确保环境 ——————
ensure_env()

import os, re, sys, unicodedata, string
from pathlib import Path
from typing import List, Dict, Set

from datetime import datetime
import random

import itertools

import json
from openai import OpenAI

import pandas as pd
from tqdm import tqdm
from sqlalchemy import create_engine, text
from rapidfuzz import fuzz, process
import numpy as np
import torch
from sentence_transformers import SentenceTransformer

try:
    from docx import Document
    import spacy
    nlp = spacy.load("en_core_web_sm", disable=["parser", "lemmatizer"])
    model_emb = SentenceTransformer(
        "sentence-transformers/all-MiniLM-L6-v2",
        device="cuda" if torch.cuda.is_available() else "cpu"
    )
except Exception:
    cute_box(
      "缺少依赖：请运行 pip install python-docx spacy",
      "依存関係が足りません：pip install python-docx spacy を実行してね",
      "⚠️"
    )
    sys.exit(1)

# ---------------- 常量 ----------------
STOPWORDS = {"the","and","for","with","from","that","this","have","will","are","you","not","but","all","any","one","our","their"}
# 预设的关键词列表 (选项1)
PRESET_KEYWORDS_2025 = [
    'partner','alliance','collaborat','cooper','cooperat','join','merger','acquisiti',
    'outsourc','invest','licens','integrat','coordinat','synergiz','associat',
    'confedera','federa','union','unit','amalgamat','conglomerat','combin',
    'buyout','companion','concur','concert','comply','complement','assist',
    'takeover','accession','procure','suppl','conjoint','support','adjust',
    'adjunct','patronag','subsid','affiliat','endors'
]
# 全局使用的关键词列表 (初始为空，稍后在 configure_keywords 中赋值)
KEYWORD_ROOTS = []

# 匹配: "April 28, 2025" 或 "21 May 2025"
DATE_FINDER = re.compile(
    r'\b(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2},?\s+\d{4}|\d{1,2}\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})\b',
    re.IGNORECASE
)
# ---------------- Bad-Rate 规则 ----------------
ORG_SUFFIX  = re.compile(
    r'\b(Inc\.?|Corp\.?|Corporation|Ltd\.?|LLC|PLC|AG|NV|SA|GmbH|S\.p\.A|Co\.?|Company|'
    r'Group|Holdings?|Partners?|Capital|Ventures?|Bank|Trust|Software|'
    r'Technolog(?:y|ies)|Pharma(?:ceuticals)?|Systems?|Services?|'
    r'Industr(?:y|ies)|Foundation|Laborator(?:y|ies)|'
    r'University|College|Institute|School|Hospital|Center|Centre)\b',
    re.I)

TIME_QTY    = re.compile(
    r'\b(year|month|week|day|decade|centur(?:y|ies)|quarter|q[1-4]|'
    r'ago|last|next|few|couple|several|dozen|half|around|approximately)s?\b',
    re.I)

# ── 金融报表 / 业绩公告类 ─────────────────────────────
FIN_REPORT = re.compile(
    r'\b(results?|earnings?|revenues?|turnover|profit(?:s)?|loss(?:es)?|guidance|forecast|'
    r'financial statements?|balance sheets?|cash flows?|income statements?)\b',
    re.I)

# ── 分季 / 分半期 / 分年描述 ─────────────────────────
ORDINAL_PERIOD = re.compile(
    r'\b(first|second|third|fourth|fifth|sixth|seventh|eighth|ninth|tenth)\b.*?\b(quarter|half|year)\b',
    re.I)

# ── 典型“公告/报告/更新”触发词（多见于 ban_list） ─────────────
ANNOUNCE_VERB = re.compile(
    r'\b(reports?|announces?|updates?|revises?|publishes?|files?|issues?|unveils?)\b',
    re.I)
# ===== 在常量区（TIME_QTY 之后）新增几条通用 regex =====
GENERIC_NOUN = re.compile(
    r'\b(services?|solutions?|systems?|platforms?|programs?|projects?|'
    r'statements?|reports?|targets?|technologies?|operations?|activities|'
    r'strategies?|plans?)\b', re.I)

MONTH_NAME = re.compile(
    r'\b(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|'
    r'jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|'
    r'dec(?:ember)?)\b', re.I)

NEW_GENERIC_TIME = re.compile(
    r'\b(?:end|beginning|middle|start|first|second|third|fourth|prior|previous|'
    r'current|next)\s+(?:of\s+)?(?:the\s+)?(?:year|quarter|month|week)s?\b',
    re.I)

#  纯大写 2–4 位缩写（PBM / ESG …）
ALLCAP_SHORT = re.compile(r'^[A-Z]{2,4}$')

#  %、百万/十亿、美元符号之类
NUMERIC = re.compile(r'[%\$]\s*\d|\d[\d,\.]+\s*(?:million|billion|thousand)', re.I)


ALL_UPPER  = re.compile(r'^[A-Z]{2,}$')
ALL_LOWER  = re.compile(r'^[a-z]{4,}$')

SHORT_TOKEN = re.compile(r'^[A-Za-z]{1,4}$')
ART_LOWER   = re.compile(r'^\s*(a|an|about|approximately|the|this|that|those)\s+[a-z]')
GENERIC_END = re.compile(
    r'\b(plan|plans?|programs?|systems?|platforms?|services?|solutions?|operations?|'
    r'agreements?|strategies?|reports?|statements?)$', re.I)

def _lower_ratio(text: str) -> float:
    w = text.split()
    return sum(t[0].islower() for t in w) / len(w) if w else 0

# --- 【新增】预定义一些“非公司”的垃圾概念向量 ---
# 这些词代表了我们想过滤掉的“垃圾类型”
NOISE_CONCEPTS = [
    "financial report results",   # 财报类
    "fiscal year quarter",        # 时间类
    "forward looking statements", # 法律声明类
    "January February March",     # 月份
    "global market growth",       # 泛指市场
    "conference call webcast",    # 会议
    "operating expenses",         # 会计术语
    "agreement partnership"       # 泛指合作
]

# 预计算垃圾概念的向量（为了速度，只算一次）
# 注意：这行代码要放在 model_emb 加载之后
print("⏳ 正在预计算垃圾词向量...")
noise_vecs = model_emb.encode(NOISE_CONCEPTS, normalize_embeddings=True)

def calc_Bad_Score(text: str) -> int:
    """
    【升级版】规则 + AI 混合评分
    """
    score = 0
    
    # === ① 规则判断 (保留原有的快速筛选，速度极快) ===
    if ORG_SUFFIX.search(text): return 0       # 像公司名，直接放行
    
    # 基础规则扣分
    if TIME_QTY.search(text): score += 30
    if FIN_REPORT.search(text): score += 30
    if len(text.split()) <= 2: score += 10
    if _lower_ratio(text) > 0.30: score += 10
    
    # === ② AI 语义判断 (新增核心功能) ===
    # 只有当 text 比较长，或者规则没判 0 分时，才动用 AI (省算力)
    if score > 0 or len(text.split()) > 2:
        # 1. 计算当前词的向量
        text_vec = model_emb.encode([text], normalize_embeddings=True)[0]
        
        # 2. 计算它和“垃圾概念”的最大相似度
        # np.dot 计算点积 (因为已经normalize了，所以等同于余弦相似度)
        sims = np.dot(noise_vecs, text_vec)
        max_sim = float(np.max(sims))
        
        # 3. 根据相似度扣分
        if max_sim > 0.4:  # 稍微有点像垃圾
            score += 20
        if max_sim > 0.6:  # 很像垃圾
            score += 40
        if max_sim > 0.8:  # 几乎确定是垃圾
            score += 100

    return score
# ---------------- 全局变量 ----------------
BASE_DIR = Path(__file__).resolve().parent
MAX_COMP_COLS = 50
SENTENCE_RECORDS: List[Dict] = []

# ---------------- 共用 ----------------



# ------- 新版本：首次输入后写 .db_key，后续自动读取 -------
def ask_mysql_url() -> str:
    key_file = Path(__file__).with_name(".db_key")   # 脚本同目录 .db_key
    if key_file.exists():
        key = key_file.read_text().strip()
    else:
        key = input("请输入秘钥/キーを入力してください：user:pass@host\n>>>>>> ").strip()
        key_file.write_text(key)                     # 缓存下次用
    return f"mysql+pymysql://{key}:3306/na_data?charset=utf8mb4" 

def choose() -> str:
    """
    显示主菜单，返回用户选择
    """
    cute_box(
        "CorpLink-AI 自动化处理系统\n"
        "------------------------------------------------\n"
        "① [开始] 提取数据 (Step 1-2)\n"
        "   - 从文档提取句子 -> 初步识别 -> 生成待清洗表\n\n"
        "② [清洗] AI 自动名寄せ (Step 2.5)\n"
        "   - 调用 GPT API 自动清洗/标准化 result_mapping_todo.csv\n\n"
        "③ [完成] 入库与分析 (Step 3-4)\n"
        "   - 读取清洗后的表 -> 存入数据库 -> 生成网络分析表\n"
        "------------------------------------------------\n"
        "作者：杨天乐 @ 关西大学 伊佐田研究室",
        
        "CorpLink-AI 自動化処理システム\n"
        "------------------------------------------------\n"
        "① [開始] データ抽出・一次処理 (Step 1-2)\n"
        "   - ドキュメント解析 -> 企業名抽出 -> 候補リスト生成\n\n"
        "② [浄化] AIによる自動名寄せ (Step 2.5)\n"
        "   - GPT APIを利用して、表記ゆれやノイズを自動修正\n\n"
        "③ [完了] DB登録・ネットワーク分析 (Step 3-4)\n"
        "   - クリーニング済みデータをDBへ登録 -> 分析用テーブル出力\n"
        "------------------------------------------------\n"
        "作成者：楊 天楽　協力：李 宗昊 李 佳璇 @関西大学",
        
        "🤖"
    )
    
    while True:
        c = input("👉 请输入功能序号 / 番号を入力してください (1/2/3): ").strip()
        if c in {"1", "2", "3"}:
            return c
        print("❌ 输入无效，请重新输入 / 無効な入力です")

# ---------- 【新增功能】关键词配置函数 ----------
# 全局变量
KEYWORD_ROOTS = []
USE_SEMANTIC_FILTER = False  # 新增：标记是否使用语义筛选

# 语义筛选的“标杆句子”
ANCHOR_TEXT = "Companies announce strategic partnership, joint venture, merger, acquisition, investment, or business collaboration."

def configure_keywords():
    """
    配置筛选模式：预设关键词、自定义关键词、或 AI 语义筛选
    """
    global KEYWORD_ROOTS, USE_SEMANTIC_FILTER
    
    cute_box(
        "【配置】请选择信息抽取的模式：\n"
        "1. 关键词模式: 2025 AI x Healthcare (默认)\n"
        "2. 关键词模式: 自定义输入\n"
        "3. AI语义模式: 语义向量匹配 (Beta)(sentence-transformers/all-MiniLM-L6-v2)",
        
        "【設定】情報抽出モードを選択してください：\n"
        "1. キーワードモード: 2025 AI x ヘルスケア (デフォルト)\n"
        "2. キーワードモード: カスタム入力 (その他)\n"
        "3. AIモード: ベクトル類似度マッチング (Beta)(sentence-transformers/all-MiniLM-L6-v2)",
        "⚙️"
    )
    
    choice = input("👉 请输入 / 番号を入力 (1/2/3) [Default: 1]: ").strip()
    
    if choice == "3":
        USE_SEMANTIC_FILTER = True
        print("\n✅ [System] AI语义筛选已启用 (Model: sentence-transformers/all-MiniLM-L6-v2)")
        print("   [System] AIフィルタリングが有効になりました")
        
    elif choice == "2":
        print("\n👉 请输入自定义关键词 (逗号分隔) / カスタムキーワードを入力 (カンマ区切り):")
        raw_input = input(">>>>>> ").strip()
        try:
            custom_keys = [k.strip().strip("'").strip('"') for k in raw_input.split(',') if k.strip()]
            if not custom_keys: raise ValueError
            KEYWORD_ROOTS = custom_keys
            print(f"✅ [System] 已加载 {len(KEYWORD_ROOTS)} 个自定义关键词")
        except:
            print("❌ [Error] 格式错误，已回退到默认模式 / フォーマットエラー、デフォルトに戻ります")
            KEYWORD_ROOTS = PRESET_KEYWORDS_2025
    else:
        KEYWORD_ROOTS = PRESET_KEYWORDS_2025
        print("✅ [System] 已加载默认关键词集 / デフォルトキーワードをロードしました")

def dedup_company_cols(df: pd.DataFrame) -> pd.DataFrame:
    comp_cols = [c for c in df.columns if c.startswith("company_")]
    for ridx in df.index:
        seen: Set[str] = set()
        for col in comp_cols:
            val = str(df.at[ridx, col]).strip()
            if val in seen:
                df.at[ridx, col] = ""
            else:
                seen.add(val)
    return df

# ---------------- Step‑1 ----------------

def _normalize(text: str) -> str:
    t = re.sub(r"\s+", " ", text.lower().strip())
    t = re.sub(r"^[\-:\"']+|[\-:\"']+$", "", t)
    t = re.sub(r"[,.;/()]+", "", t)
    return t.strip()

def clean_text(t: str) -> str:
    return ''.join(c for c in t if unicodedata.category(c)[0] != 'C' or c in ('\n', '\t'))

def extract_sentences(path: Path) -> List[str]:
    doc = Document(path)
    collecting, current, articles = False, "", []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt: continue
        tag = txt.lower()
        if tag == "body": collecting, current = True, ""; continue
        if tag in ("notes", "classification") and collecting:
            collecting = False; articles.append(current.strip()); continue
        if collecting: current += " " + txt
    sents = []
    for art in articles:
        for s in re.split(r"\.\s*", art):
            s = s.strip();
            if len(s) >= 5: sents.append(s)
    return sents

def extract_index_titles(paragraphs):
    paras_text = [p.text.strip() for p in paragraphs]
    
    # 1. 尝试获取总篇数
    m = re.search(r'Documents?\s*\(\s*(\d+)\s*\)', '\n'.join(paras_text[:50]), re.I)
    if not m: return []
    total = int(m.group(1))
    
    pat = re.compile(r'^(\d+)\.\s+(.*)$')
    seen = set()
    titles = []
    
    for i, line in enumerate(paras_text):
        m2 = pat.match(line)
        if m2:
            # === 【核心修复】验证机制 ===
            # 真正的目录标题，后面 6 行内一定有 "Client/Matter"
            is_valid_toc = False
            for offset in range(1, 7): 
                if i + offset >= len(paras_text): break
                next_line = paras_text[i + offset].lower()
                if "client/matter" in next_line or "search terms" in next_line:
                    is_valid_toc = True
                    break
            
            # 如果后面没有元数据，说明这是正文里的普通列表（如 "3. REVENUE"），跳过！
            if not is_valid_toc:
                continue 
            
            # === 验证通过 ===
            raw = m2.group(2).strip()
            norm = _normalize(raw)
            
            if norm in seen: continue
            seen.add(norm)
            titles.append((int(m2.group(1)), raw, norm))
            
            if len(titles) >= total: break
                
    return sorted(titles, key=lambda x: x[0])

def extract_sentences_by_titles(filepath: str) -> List[Dict]:
    """
    【修复版 v4】
    1. 包含日期验证机制：防止匹配到正文里的重复标题，解决 Publisher 抓错问题。
    2. 包含语义筛选功能：支持关键词/AI语义模式切换。
    """
    doc = Document(filepath); paras = doc.paragraphs
    index_titles = extract_index_titles(paras); recs = []
    
    if index_titles:
        paras_norm = [_normalize(p.text) for p in paras]
        last_article_end_idx = 0

        for i_title, (doc_idx, title_raw, title_norm) in enumerate(index_titles):
            match_idx = -1
            date_line_idx = -1
            
            # 从上次结束位置开始，寻找所有匹配的标题行
            candidates = [i for i, n in enumerate(paras_norm) 
                          if i >= last_article_end_idx and n == title_norm]
            
            for idx in candidates:
                # 1. 目录检查 (Client/Matter) - 跳过目录
                if idx + 1 < len(paras):
                    next_line = paras[idx+1].text.strip().lower()
                    if next_line.startswith("client/matter") or next_line.startswith("search terms"):
                        continue 

                # 2. 【关键修复】日期验证机制
                # 真正的文章标题，其后 1-3 行内一定包含日期。
                # 如果找不到日期，说明这是正文里的假标题（摘要或引用），跳过！
                found_date = False
                temp_date_idx = -1
                
                # 往后看 3 行寻找日期
                for offset in range(1, 4):
                    if idx + offset >= len(paras): break
                    txt = paras[idx + offset].text.strip()
                    if DATE_FINDER.search(txt):
                        found_date = True
                        temp_date_idx = idx + offset
                        break
                
                if found_date:
                    match_idx = idx
                    date_line_idx = temp_date_idx
                    break
            
            # 如果没找到带日期的标题，就放弃这篇文章（防止抓错）
            if match_idx == -1: 
                continue

            # ----------------------------------

            # 3. 智能提取 Publisher
            # Publisher 通常在 Title 和 Date 之间
            # 如果 Date 在 Title 下面第 2 行或更远，中间那就是 Publisher
            if date_line_idx > match_idx + 1:
                publisher = paras[match_idx + 1].text.strip()
            else:
                publisher = "" # 只有日期，没写出版社
            
            # 4. 提取日期
            news_date = ""
            m = DATE_FINDER.search(paras[date_line_idx].text.strip())
            if m: news_date = m.group(0)

            # 5. 确定正文范围
            # Body 应该在 Date 之后开始，更新 pub_idx 为日期行，方便后续定位
            pub_idx = date_line_idx 
            
            search_end_limit = len(paras)
            if i_title + 1 < len(index_titles):
                next_title_norm = index_titles[i_title+1][2]
                try:
                    next_candidates = [i for i, n in enumerate(paras_norm) 
                                       if i > match_idx + 20 and n == next_title_norm]
                    if next_candidates: search_end_limit = next_candidates[0]
                except Exception: pass

            body_start = next((i+1 for i in range(pub_idx+1, search_end_limit) if paras[i].text.strip().lower() == "body"), None)
            if body_start is None: 
                body_start = pub_idx + 1 # 如果没 Body 标签，就从日期下一行开始
            
            body_end = len(paras)
            for i in range(body_start, search_end_limit):
                t_low = paras[i].text.strip().lower()
                if t_low.startswith("notes") or t_low.startswith("classification") or "(end) dow jones" in t_low:
                    body_end = i
                    break
            last_article_end_idx = body_end

            # 6. 提取句子 (保留了语义筛选逻辑)
            article = " ".join(clean_text(paras[i].text) for i in range(body_start, body_end))
            
            # 切分句子
            raw_sents = [s.strip() for s in re.split(r"\.\s*", article) if len(s.strip())>=20]

            # --- 如果是语义模式，先批量计算向量 ---
            if 'USE_SEMANTIC_FILTER' in globals() and USE_SEMANTIC_FILTER and raw_sents:
                # 计算标杆向量 (如果还没算过)
                if not hasattr(extract_sentences_by_titles, "anchor_vec"):
                     extract_sentences_by_titles.anchor_vec = model_emb.encode([ANCHOR_TEXT], normalize_embeddings=True)[0]
                
                sent_vecs = model_emb.encode(raw_sents, normalize_embeddings=True)
                sim_scores = np.dot(sent_vecs, extract_sentences_by_titles.anchor_vec)
            else:
                sim_scores = [0.0] * len(raw_sents)

            # --- 遍历句子进行筛选 ---
            for i, sent in enumerate(raw_sents):
                is_hit = False
                match_reason = ""
                hit_count = 0

                if 'USE_SEMANTIC_FILTER' in globals() and USE_SEMANTIC_FILTER:
                    # === 模式 A: AI 语义筛选 ===
                    score = float(sim_scores[i])
                    if score > 0.45: # 阈值可调
                        is_hit = True
                        hit_count = 1 # 语义命中算 1 分
                        match_reason = f"Semantic({score:.2f})"
                else:
                    # === 模式 B: 关键词筛选 ===
                    hits = [k for k in KEYWORD_ROOTS if k in sent.lower()]
                    if hits:
                        is_hit = True
                        hit_count = len(hits)
                        match_reason = "; ".join(hits)
                
                if is_hit:
                    recs.append({
                        "Title": title_raw,
                        "Publisher": publisher,
                        "Date": news_date,
                        "Country": "",
                        "Sentence": sent,
                        "Hit_Count": hit_count,
                        "Matched_Keywords": match_reason
                    })
        
        if recs: return recs

    # Fallback (无索引情况)
    # 这里不需要太复杂的日期验证，因为单篇文章通常结构简单
    for sent in extract_sentences(Path(filepath)):
        # 这里只做了简单的关键词筛选兼容，如果需要Fallback也支持语义，逻辑同上
        hits = [k for k in KEYWORD_ROOTS if k in sent.lower()]
        if hits:
             recs.append({
                "Title": "", "Publisher": "", "Date": "", "Country": "", 
                "Sentence": sent, "Hit_Count": len(hits), 
                "Matched_Keywords": "; ".join(hits)
            })
    return recs

def step1():
    cute_box(
        "Step-1：提取 Word 句子 中…",
        "Step-1：文抽出中…",
        "📄"
    )
    all_recs: List[Dict] = []

    # 1) 收集所有 .docx 路径
    docx_files = []
    for root, _, files in os.walk(BASE_DIR):
        for fname in files:
            if not fname.endswith(".docx") or fname.startswith("~$"):
                continue
            full = Path(root) / fname
            rel = full.relative_to(BASE_DIR).parts
            tier1 = rel[0] if len(rel) >= 1 else ""
            tier2 = rel[1] if len(rel) >= 2 else ""
            docx_files.append((str(full), tier1, tier2, fname))

    # 2) 逐文件提取句子
    for fp, t1, t2, fname in tqdm(docx_files, desc="📄 处理 Word 文件"):
        for r in extract_sentences_by_titles(fp):
            if not r["Title"]:
                r["Title"] = Path(fname).stem
            r.update({"Tier_1": t1, "Tier_2": t2, "Filename": fname})
            all_recs.append(r)

    global SENTENCE_RECORDS
    SENTENCE_RECORDS = all_recs
    cute_box(
        f"Step-1 完成，共 {len(all_recs)} 条记录",
        f"Step-1 完了しました：全{len(all_recs)}件",
        "✅"
    )

# ----------------—— Step‑2 ——----------------

def is_valid_token(token: str) -> bool:
    token = token.strip()
    if "@" in token or token.startswith("http"):    # ① 含邮箱 / URL 特征
        return False
    if not token or all(c in "-–—・.、。！？／ー" for c in token):
        return False
    if re.search(r"\d", token) and not re.search(r"[A-Za-z]", token):
        return False
    if "  " in token:
        return False
    return True


# —— 4. 原始企业名提取 ——
# —— 4. 原始企业名提取 ——
def extract_companies(text: str,
                      company_db: List[str],
                      ner_model,
                      fuzzy_threshold: int = 95) -> List[str]:
    """
    · **仅负责“把句子里可能是公司名的片段全部抓出来”**，
      不做任何 ban/映射/去重处理——这些留给后续数据库比对阶段完成。
    · 识别逻辑完全沿用单体版（spaCy NER + “IBMers”正则 + 严格模糊匹配）。
    """
    comps: Set[str] = set()

    # 1) 去掉日期（排除 ‘xx/xx/xxxx 之后整段’ 的噪音）
    text_clean = re.sub(r"\s*\d{1,2}/\d{1,2}/\d{2,4}.*$", "", text).strip()
    # --- 新增清洗 ---
    # 1) 去掉 ® ™ ©
    text_clean = re.sub(r"[®™©]", "", text_clean)
    # 2) 去掉简写商标括号，如 “Weight Doctors(R)”
    text_clean = re.sub(r"\(\s*[A-Z]{1,3}\s*\)", "", text_clean)
    # 3) 整句里带邮箱的直接剪掉邮箱
    text_clean = re.sub(r"\b\S+@\S+\b", "", text_clean)

    # 2) spaCy NER
    doc = ner_model(text_clean)
    for ent in doc.ents:
        ent_text = ent.text.strip()

        # —— 基础噪音过滤（和单体版一致）
        if "  " in ent_text or re.search(r"[\d/%+]|[^\x00-\x7F]", ent_text):
            continue
        valid_ent = True
        for w in ent_text.split():
            if (not w[0].isalpha()
                or w in {"The","And","For","With","From","That","This"}
                or not is_valid_token(w)):
                valid_ent = False
                break
        if valid_ent:
            comps.add(ent_text)

    # 3) “IBMers” 一类写法
    for m in re.findall(r"\b([A-Z]{2,})ers\b", text_clean):
        comps.add(m)

    # 4) 仅用于“确认是已知公司”，但依旧返回原词
    STOPWORDS = {"The","And","For","With","From","That","This","Have","Will",
                "Are","You","Not","But","All","Any","One","Our","Their"}

    tokens = re.findall(r"\b\S+\b", text_clean)
    for pos, token in enumerate(tokens):
        # —— 噪音与格式过滤（同原先逻辑） ——
        if (pos == 0 or token in STOPWORDS
            or any(ch in token for ch in "/%+") or "  " in token
            or len(token) < 5 or not token[0].isupper() or token.isupper()
            or re.search(r"\d|[^\x00-\x7F]", token)
            or not is_valid_token(token)):
            continue

        # 若数据库里存在“完全同名（大小写不同视同）”的条目，就保留；否则忽略
        if any(token.lower() == db.lower() for db in company_db):
            comps.add(token)

    return list(comps)


def step2(mysql_url: str):
    cute_box(
        "Step-2：公司识别＋BAN 过滤 中…",
        "Step-2：企業名認識＋BAN フィルタ中…",
        "🏷️"
    )
    # 单独导出 canonical 表（engine_tmp）
    engine_tmp = create_engine(mysql_url)            # ← 新建
    df_canon = pd.read_sql("SELECT id, canonical_name FROM company_canonical", engine_tmp)
    df_canon.to_csv(BASE_DIR / "canonical_list.csv", index=False, encoding="utf-8-sig")
    cute_box(
        f"已写出 canonical_list.csv，共 {len(df_canon)} 行",
        f"canonical_list.csv を保存しました：{len(df_canon)} 行",
        "🗂️"
    )
    # ---- 连接数据库 ----
    engine = create_engine(mysql_url)
    with engine.begin() as conn:
        ban_set = {r[0] for r in conn.execute(text("SELECT alias FROM ban_list"))}
        rows = conn.execute(text("""
            SELECT a.alias, c.canonical_name FROM company_alias a
            JOIN company_canonical c ON a.canonical_id = c.id
        """))
        alias_map = {alias: canon for alias, canon in rows}
        canon_set = {r[0] for r in conn.execute(text("SELECT canonical_name FROM company_canonical"))}
        # —— 预编码全部 canonical，一次搞定 ——
        canon_names = list(canon_set)
        canon_vecs  = model_emb.encode(canon_names, batch_size=64, normalize_embeddings=True)
        # ↓↓↓ 新增：名字→ID 的字典，用于 Advice 对应的 ID
        rows2 = conn.execute(text(
            "SELECT id, canonical_name FROM company_canonical"
        ))
        canon_name2id = {name: cid for cid, name in rows2}      # ← 新增
    
    cute_box(
    f"ban_list={len(ban_set)}，alias_map={len(alias_map)}，canon_set={len(canon_set)}",
    f"ban_list：{len(ban_set)}件／alias_map：{len(alias_map)}件／canon_set：{len(canon_set)}件",
    "🔍"
    )

    df = pd.DataFrame(SENTENCE_RECORDS)
    df_hit = df[df["Hit_Count"].astype(int) >= 1].reset_index(drop=True)
    if df_hit.empty:
        cute_box(
        "Step-1 没提取到任何句子，请先跑 Step-1！",
        "Step-1 で文が取得できませんでした。まず Step-1 を実行してね",
        "🚫"
        )
        return

    company_db = list(canon_set) + list(alias_map.keys())   # canonical + alias
    comp_cols: List[List[str]] = []
    for sent in tqdm(df_hit["Sentence"].tolist(), desc="公司识别"):
        names_raw = extract_companies(sent, company_db, nlp)
        uniq: List[str] = []
        for alias in names_raw:
            if alias in uniq:     
                continue
            uniq.append(alias)                         # 保留句面原词，不做任何替换
        comp_cols.append(uniq[:MAX_COMP_COLS])

    for i in range(MAX_COMP_COLS):
        df_hit[f"company_{i+1}"] = [lst[i] if i < len(lst) else "" for lst in comp_cols]
        
# === ③ 先按数据库规则处理每行 company_n 列 ===
    ban_lower     = {b.lower() for b in ban_set}
    canon_lower   = {c.lower() for c in canon_set}
    alias_lower   = {a.lower(): canon for a, canon in alias_map.items()}
    canon_lower2orig = {c.lower(): c for c in canon_set}
    # —— 标准化：去掉所有非字母数字，再小写
    def _norm_key(s: str) -> str:
        return re.sub(r"[^A-Za-z0-9]", "", s).lower()

    comp_cols = [f"company_{i+1}" for i in range(MAX_COMP_COLS)]

    for ridx in df_hit.index:
        orig_names = [df_hit.at[ridx, c].strip() for c in comp_cols if df_hit.at[ridx, c].strip()]
        new_names  = []
        for nm in orig_names:
            nm_l = nm.lower()
            # ① ban → 丢弃
            if nm_l in ban_lower:
                continue
            # ② 已是标准名 → 保留原样
            if nm_l in canon_lower:
                new_names.append(canon_lower2orig[nm_l])
                continue
            # ③ 别名 → 替换为对应 canonical
            if nm_l in alias_lower:
                new_names.append(alias_lower[nm_l])
                continue
            # ④ 未知 → 原样
            new_names.append(nm)

        # ⑤ 顺位左移 + “同根” 去重
        cleaned = []
        seen_keys = set()
        for nm in sorted(new_names, key=len, reverse=True):           # 先长后短
            key = _norm_key(nm)
            # 1) 与已选任何名称 key 前缀 / 后缀 相同 → 视为重复
            if any(key in k or k in key for k in seen_keys):
                continue
            cleaned.append(nm)
            seen_keys.add(key)
        # ⑥ 写回行（不足补空，用 .at）
        for i, col in enumerate(comp_cols):
            df_hit.at[ridx, col] = cleaned[i] if i < len(cleaned) else ""


    # === ④ 继续原流程写 result.csv（下方原代码保持不变） ===

    # ---- 组装精简版 result.csv ----
    meta_cols = ["Tier_1", "Tier_2", "Filename", "Date",
                 "Title", "Publisher", "Sentence",
                 "Hit_Count", "Matched_Keywords"]

    df_final = (df_hit[meta_cols +
                [c for c in df_hit.columns if c.startswith("company_")]]
                .fillna(""))
    df_final = dedup_company_cols(df_final)

    df_final.to_csv(BASE_DIR / "result.csv",
                    index=False, encoding="utf-8-sig")
    cute_box(
        f"已生成 result.csv，共 {len(df_final)} 条记录",
        f"result.csv を生成しました：全{len(df_final)}件",
        "📑"
    )
    
    # ---- 生成 result_mapping_todo.csv （空表安全 + 统计）----
    # 1) name→id 字典（Advice 需要）
    canon_name2id = {row.canonical_name: row.id for row in df_canon.itertuples()}

    todo_rows: List[Dict] = []

    # 统计：哪些被过滤/跳过
    ban_hits = alias_hits = canon_hits = 0
    rows_skipped_not_enough_companies = 0  # 同行公司总数 < 2 的行

    comp_cols = [c for c in df_final.columns if c.startswith("company_")]

    for _, row in df_final.iterrows():
        # 取出该行所有非空企业名（已做过ban/alias/canonical一次清洗与同根去重）
        names = [row[c].strip() for c in comp_cols if row[c].strip()]

        # —— 分类统计（先更新 ban/alias/canonical 计数，再收集 unknowns）
        unknowns: List[str] = []
        for alias in names:
            alias_l = alias.lower()
            if alias_l in ban_lower:
                ban_hits += 1
                continue
            if alias_l in alias_lower:
                alias_hits += 1
                continue
            if alias_l in canon_lower:
                canon_hits += 1
                continue
            unknowns.append(alias)
        # ✅ 行级门槛：同行公司总数（names）< 2 → 不入 todo，但上面的统计已计入
        if len(names) < 2:
            rows_skipped_not_enough_companies += 1
            continue

        # 该行达到“有2+家公司”的门槛，即使 unknowns 只有1个或更多，都进入 todo
        # ---------------- 优化后的匹配逻辑 (Fuzzy + AI) ----------------
        
        # 1. 批量计算本行 unknowns 的向量 (性能优化：一次算完，比循环里一个个算快得多)
        if len(canon_vecs) > 0 and unknowns:
            unknown_vecs = model_emb.encode(unknowns, normalize_embeddings=True)
        else:
            unknown_vecs = []

        for i, alias in enumerate(unknowns):
            advice = ""
            adviced_id = ""
            match_info = "" # 调试用，看看是谁立功了
            
            # --- 策略 A: 高精度模糊匹配 (RapidFuzz) ---
            # 专治：拼写错误、后缀差异 (e.g. "Apple Incc" vs "Apple Inc.")
            # token_sort_ratio 可以忽略单词顺序 (e.g. "Motors General" vs "General Motors")
            fuzzy_res = process.extractOne(alias, canon_names, scorer=fuzz.token_sort_ratio)
            
            if fuzzy_res:
                candidate, score, _ = fuzzy_res
                # 设定一个较高的门槛，比如 90 分，确保字面非常像才直接采纳
                if score >= 90:
                    advice = candidate
                    adviced_id = canon_name2id.get(advice, "")
                    match_info = f"Fuzzy({score:.0f})"
            
            # --- 策略 B: AI 向量语义匹配 ---
            # 只有当 Fuzzy 没搞定 (advice 为空) 时，才请 AI 出山
            if not advice and len(canon_vecs) > 0:
                # 取出刚刚批量算好的向量
                curr_vec = unknown_vecs[i]
                
                # 计算与所有标准名的相似度
                sims = np.dot(canon_vecs, curr_vec)
                best_idx = int(np.argmax(sims))
                vector_score = float(sims[best_idx])
                
                # 阈值：0.82 (稍微提高一点门槛，减少幻觉)
                if vector_score >= 0.82:
                    advice = canon_names[best_idx]
                    adviced_id = canon_name2id.get(advice, "")
                    match_info = f"AI({vector_score:.2f})"

            # --- 存入 Todo ---
            todo_rows.append({
                "Sentence": row["Sentence"],
                "Alias":    alias,
                "Bad_Score": calc_Bad_Score(alias),
                "Advice":   advice,           # 推荐结果
                "Adviced_ID": adviced_id,     # 推荐ID
                # "Match_Info": match_info,   # (可选) 如果你想在 CSV 里看到是 AI 还是 Fuzzy 匹配的，可以把这行加到 csv 列里
                "Canonical_Name": "",
                "Std_Result": ""
            })

    # 2) 组装 DataFrame（空表安全）
    todo_cols = [
        "Sentence", "Alias", "Bad_Score",
        "Advice", "Adviced_ID",
        "Canonical_Name", "Std_Result"
    ]

    if not todo_rows:
        # —— 没有新的别名需要映射：写出只有表头的空表，并友好提示
        todo_df = pd.DataFrame(columns=todo_cols)
        todo_df.to_csv(BASE_DIR / "result_mapping_todo.csv",
                       index=False, encoding="utf-8-sig")

        cute_box(
            "本批没有产生新的别名需要映射；已被规则识别/过滤，或因“同行公司不足（<2）”规则而跳过。\n"
            f"（ban 命中：{ban_hits}，已有 alias：{alias_hits}，已有 canonical：{canon_hits}，同行公司不足跳过：{rows_skipped_not_enough_companies}）",
            "今回のバッチでは新しい別名はありません。既存データに一致／除外、または「同一行の企業数が2未満」規則でスキップされました。\n"
            f"ban 一致：{ban_hits}／既存エイリアス：{alias_hits}／既存カノニカル：{canon_hits}／同一行の企業数不足スキップ：{rows_skipped_not_enough_companies}",
            "ℹ️"
        )
    else:
        # —— 正常去重（按别名小写）
        todo_df = pd.DataFrame(todo_rows)
        todo_df["__alias_l"] = todo_df["Alias"].str.lower()
        todo_df = todo_df.drop_duplicates("__alias_l").drop(columns="__alias_l")

        # 分组排序
        todo_df["__grp"] = todo_df["Bad_Score"].apply(lambda x: 0 if x >= 50 else (1 if x >= 10 else 2))
        todo_df = (todo_df
                   .sort_values(["__grp", "Sentence"], ascending=[True, True])
                   .drop(columns="__grp"))

        # 固定列顺序
        for col in todo_cols:
            if col not in todo_df.columns:
                todo_df[col] = ""   # 兜底，保证列齐全
        todo_df = todo_df[todo_cols]

        # 写文件
        todo_df["Bad_Score"] = todo_df["Bad_Score"].astype(int).astype(str)
        todo_df['Sentence'] = todo_df['Sentence'].apply(
            lambda s: "'" + s if isinstance(s, str) and s.startswith('=') else s
        )
        todo_df.to_csv(BASE_DIR / "result_mapping_todo.csv",
                       index=False, encoding="utf-8-sig")

        cute_box(
            f"已生成 result_mapping_todo.csv，共 {len(todo_df)} 条待处理别名。\n"
            f"（ban 命中：{ban_hits}，已有 alias：{alias_hits}，已有 canonical：{canon_hits}，同行公司不足跳过：{rows_skipped_not_enough_companies}）",
            f"result_mapping_todo.csv を作成：{len(todo_df)} 件の候補。\n"
            f"（ban 一致：{ban_hits}／既存エイリアス：{alias_hits}／既存カノニカル：{canon_hits}／同一行の企業数不足スキップ：{rows_skipped_not_enough_companies}）",
            "📝"
        )
        
    cute_box(
    "Step-2 完成！请编辑 result_mapping_todo.csv 然后运行 Step-3",
    "Step-2 完了！result_mapping_todo.csv を編集してから Step-3 を実行してね",
    "✅"
    )
    cute_box(
        "result_mapping_todo.csv 快速填写指南：\n"
        "1) 空白→跳过\n"
        "2) 0→加 ban_list\n"
        "3) n→视为 canonical_id\n"
        "4) 其他→新/已有标准名",
        "result_mapping_todo.csv 簡易入力ガイド：\n"
        "1) ブランク→スキップ\n"
        "2) 0→ban_list登録\n"
        "3) n→canonical_id と見なす\n"
        "4) その他→新規/既存標準名",
        "📋"
    )

# ================ Step-3 ==============

def step3(mysql_url: str):
    """
    Step-3 标准化 + 写库（与旧 NA_step3_standardize.py 等价）
    - Canonical_Name == ''  → Std_Result = 'No input'
    - Canonical_Name == '0' → 写 ban_list,  Std_Result = 'Banned'
    - 其它:
        • 若已存在 alias → Std_Result = 'Exists'
        • 否则插入/补全 canonical & alias, Std_Result = 'Added'
    同时把最新映射应用回 result.csv
    """
    # 本轮批次号：YYYYMMDD + 8位随机数
    process_id = datetime.now().strftime("%Y%m%d") + f"{random.randint(0, 99999999):08d}"
    res_f  = BASE_DIR / "result.csv"
    todo_f = BASE_DIR / "result_mapping_todo.csv"
    if not (res_f.exists() and todo_f.exists()):
        cute_box(
            "找不到 result.csv 或 result_mapping_todo.csv，请先生成它们",
            "result.csv または result_mapping_todo.csv が見つかりません。先に作成してね",
            "❗"
        )
        sys.exit(1)

    # 读取现有文件
    df_res = pd.read_csv(res_f,  dtype=str).fillna("")
    df_map = pd.read_csv(todo_f, dtype=str).fillna("")
    if "Process_ID" not in df_map.columns:
        df_map["Process_ID"] = ""

    engine = create_engine(mysql_url)
    with engine.begin() as conn:
        # 拉取三表到内存
        ban_set    = {r[0] for r in conn.execute(text("SELECT alias FROM ban_list"))}
        canon_map  = {r[0]: r[1] for r in conn.execute(text("SELECT id, canonical_name FROM company_canonical"))}
        alias_map  = {r[0]: r[1] for r in conn.execute(text(
            "SELECT a.alias, c.canonical_name FROM company_alias a "
            "JOIN company_canonical c ON a.canonical_id=c.id"
        ))}
        # 无视大小写的镜像
        ban_lower       = {b.lower() for b in ban_set}
        alias_lower_map = {a.lower(): c for a, c in alias_map.items()}
        canon_lower2id  = {name.lower(): cid for cid, name in canon_map.items()}

    # —— 处理 todo 映射 ——  
    for idx, row in df_map.iterrows():
        alias_raw   = row["Alias"].strip()
        alias_raw_l = alias_raw.lower()
        canon_input = row["Canonical_Name"].strip()
        if not canon_input:
            df_map.at[idx, "Std_Result"] = "No input"
            continue

        # Case A: Ban (0)
        if canon_input == "0":
            if alias_raw_l not in ban_lower:
                try:
                    with engine.begin() as conn:
                        conn.execute(text(
                            "INSERT IGNORE INTO ban_list(alias, process_id) VALUES (:a, :pid)"
                        ), {"a": alias_raw, "pid": process_id})
                    ban_lower.add(alias_raw_l)
                except Exception as e:
                    print(f"⚠️ Ban insert skip: {e}")
            df_map.at[idx, "Std_Result"]   = "Banned"
            df_map.at[idx, "Process_ID"] = f"'{process_id}"
            continue

        # Case B: Existing ID (数字)
        if canon_input.isdigit():
            cid = int(canon_input)
            if cid not in canon_map:
                df_map.at[idx, "Std_Result"] = "Bad ID"
                continue
            canon_name = canon_map[cid]
            
        # Case C: New/Text Canonical
        else:
            ci_l = canon_input.lower()
            # 如果内存里没这个公司，尝试插入
            if ci_l not in canon_lower2id:
                try:
                    # --- 尝试插入新公司 ---
                    with engine.begin() as conn:
                        res = conn.execute(text(
                            "INSERT INTO company_canonical(canonical_name, process_id) VALUES (:c, :pid)"
                        ), {"c": canon_input, "pid": process_id})
                    new_id = res.lastrowid
                    
                except Exception as e:
                    # --- 【关键修复】如果报错(重复)，说明数据库里其实已经有了 ---
                    # 可能是因为重音符号(É vs E)导致 Python 没认出来，但数据库认出来了
                    print(f"⚠️ 发现潜在重复公司: {canon_input}，尝试从数据库获取 ID...")
                    with engine.begin() as conn:
                        # 尝试直接用名字查 ID
                        rows = conn.execute(text(
                            "SELECT id FROM company_canonical WHERE canonical_name = :c"
                        ), {"c": canon_input}).fetchall()
                        
                        if rows:
                            new_id = rows[0][0]
                            print(f"   -> 已找回现有 ID: {new_id}")
                        else:
                            # 极其罕见的情况：插入报错但又查不到，记录错误跳过
                            print(f"❌ 无法解决的冲突，跳过此条: {e}")
                            df_map.at[idx, "Std_Result"] = "DB Error"
                            continue

                # 更新内存字典
                canon_map[new_id]        = canon_input
                canon_lower2id[ci_l]     = new_id
                df_map.at[idx, "Process_ID"] = f"'{process_id}"
                canon_name = canon_input
            else:
                # 内存里已经有了，直接用
                canon_name = canon_map[canon_lower2id[ci_l]]

        # Case D: Insert Alias
        if alias_raw_l in alias_lower_map or alias_raw_l in canon_lower2id:
            df_map.at[idx, "Std_Result"] = "Exists"
            continue
            
        try:
            with engine.begin() as conn:
                conn.execute(text(
                    "INSERT IGNORE INTO company_alias(alias, canonical_id, process_id) "
                    "VALUES (:a, :cid, :pid)"
                ), {"a": alias_raw, "cid": canon_lower2id[canon_name.lower()], "pid": process_id})
            alias_lower_map[alias_raw_l] = canon_name
            df_map.at[idx, "Std_Result"]   = "Added"
            df_map.at[idx, "Process_ID"] = f"'{process_id}"
        except Exception as e:
            print(f"⚠️ Alias insert error: {e}")

    # 先写回 todo，再做回写 result.csv
    df_map.to_csv(todo_f, index=False, encoding="utf-8-sig")

    # ====== 将最新映射应用回 result.csv ======
    # 重新拉取 ban/alias/canonical 准备映射
    with engine.begin() as conn2:
        ban_set2    = {r[0] for r in conn2.execute(text("SELECT alias FROM ban_list"))}
        rows2       = conn2.execute(text(
            "SELECT a.alias, c.canonical_name FROM company_alias a "
            "JOIN company_canonical c ON a.canonical_id=c.id"
        ))
        alias_map2  = {a: c for a, c in rows2}
        canon_set2  = {r[0] for r in conn2.execute(text("SELECT canonical_name FROM company_canonical"))}

    ban_lower2        = {b.lower() for b in ban_set2}
    alias_lower_map2  = {a.lower(): c for a, c in alias_map2.items()}
    canon_lower2orig2 = {c.lower(): c for c in canon_set2}

    comp_cols = [c for c in df_res.columns if c.startswith("company_")]

    def _norm_key(s: str) -> str:
        return re.sub(r"[^A-Za-z0-9]", "", str(s)).lower()

    changed_cells = 0
    for ridx in df_res.index:
        # 读出原值
        orig = df_res.loc[ridx, comp_cols].astype(str).tolist()
        vals_in = [v.strip() for v in orig if v.strip()]
        vals_out = []
        for nm in vals_in:
            key = nm.lower()
            if key in ban_lower2:
                continue
            if key in alias_lower_map2:
                nm = alias_lower_map2[key]
                changed_cells += 1
                key = nm.lower()
            elif key in canon_lower2orig2:
                corrected = canon_lower2orig2[key]
                if corrected != nm:
                    changed_cells += 1
                nm = corrected
            vals_out.append(nm)
        # 同根去重 + 左移
        cleaned, seen = [], set()
        for nm in sorted(vals_out, key=len, reverse=True):
            k = _norm_key(nm)
            if any(k in kk or kk in k for kk in seen):
                continue
            cleaned.append(nm)
            seen.add(k)
        # 回写
        for i, col in enumerate(comp_cols):
            new_val = cleaned[i] if i < len(cleaned) else ""
            if str(df_res.at[ridx, col]) != new_val:
                changed_cells += 1
            df_res.at[ridx, col] = new_val

    # 再清一次列内重复
    df_res = dedup_company_cols(df_res)

    cute_box(
        f"已将最新映射应用到 result.csv（变更单元格约 {changed_cells} 个）",
        f"最新のマッピングを result.csv に適用しました（変更セル数 約 {changed_cells}）",
        "🛠️"
    )

    # 最后保存
    df_res.to_csv(res_f, index=False, encoding="utf-8-sig")

    cute_box(
        f"Step-3 完成，处理 {len(df_map)} 条映射，result.csv 已更新",
        f"Step-3 完了：{len(df_map)}件 処理済み，result.csv 更新完了",
        "🚀"
    )
    cute_box(
        f"本批次 Process ID：{process_id}",
        f"今回の Process ID：{process_id}",
        "📌"
    )
               
# ---------- 【新增】选项 3：GPT 自动填充功能 ----------

def ask_gpt_batch(batch_data: List[Dict], api_key: str) -> Dict:
    client = OpenAI(api_key=api_key)
    prompt = f"""
    You are a data cleaning expert for business strategy research. 
    Analyze the list of "alias" strings.

    Task: Determine the [Organizational Entity] behind the alias.
    
    [Allowed Categories] -> Set "is_company": true
    1. Commercial Companies (e.g., Toyota, Google, OpenAI)
    2. Educational Institutions (e.g., Harvard University, Tokyo High School)
    3. Government Bodies & Municipalities (e.g., Osaka Prefecture, Ministry of Economy)
    4. NGOs, NPOs, Associations (e.g., Red Cross, IEEE)
    
    [Special Mapping Rules for Products & IPs] -> Set "is_company": true
    If the 'alias' is a Product, Service, or Fictional Character/IP, DO NOT reject it. instead, map it to its OWNER Company.
    Examples:
    - "iPhone" -> is_company: true, clean_name: "Apple"
    - "ChatGPT" -> is_company: true, clean_name: "OpenAI"
    - "Mickey Mouse" -> is_company: true, clean_name: "Disney"
    - "Mario" -> is_company: true, clean_name: "Nintendo"
    - "Barbie" -> is_company: true, clean_name: "Mattel"

    [Forbidden Categories] -> Set "is_company": false
    1. General Nouns / Not Proper Nouns (e.g., "external researchers", "local governments", "our partners", "the committee", "anime", "video games")
    2. Job Titles / Departments (e.g., "CEO", "Sales Department")
    3. Individuals (unless the name refers to a sole proprietorship/studio)

    Rules for "clean_name":
    - Remove legal suffixes (Inc., Ltd., Corp., K.K., etc.).
    - If it is a Product/IP, use the OWNER Company Name.
    - Keep the full proper name (e.g., "University of Tokyo" -> "University of Tokyo").
    
    Input: {json.dumps(batch_data, ensure_ascii=False)}
    
    Output JSON format:
    {{
        "alias_original_text": {{ 
            "is_company": bool, 
            "clean_name": str, 
            "matches_advice": bool // If the mapped company matches the provided 'advice' entity
        }}
    }}
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0
        )
        return json.loads(response.choices[0].message.content)
    except:
        return {}

def step_ai_autofill():
    """
    读取 result_mapping_todo.csv，利用 GPT 自动填写 Canonical_Name 列
    包含自动保存 API Key 的功能
    """
    csv_path = BASE_DIR / "result_mapping_todo.csv"
    if not csv_path.exists():
        cute_box("找不到 result_mapping_todo.csv！", "ファイルが見つかりません", "❌")
        return

    # --- 【新增】自动读取/保存 Key 的逻辑 ---
    key_file = BASE_DIR / ".openai_key"
    api_key = ""
    
    if key_file.exists():
        api_key = key_file.read_text().strip()
        print(f"🔑 已自动加载保存的 API Key: {api_key[:8]}...")
    
    if not api_key:
        api_key = input("请输入 OpenAI API Key (sk-...) / APIキーを入力: ").strip()
        if api_key:
            # 保存到文件
            key_file.write_text(api_key)
            print("💾 API Key 已保存，下次无需输入。")
    # ----------------------------------------

    if not api_key:
        print("❌ 未输入 Key，操作取消。")
        return

    print("⏳ 正在读取 CSV...")
    df = pd.read_csv(csv_path, dtype=str).fillna("")
    
    # ... (后续逻辑保持不变) ...
    
    # 2. 筛选出需要处理的行 (Canonical_Name 为空的行)
    mask = df["Canonical_Name"] == ""
    rows_to_process = df[mask]
    
    if rows_to_process.empty:
        print("✨ 所有行的 Canonical_Name 都已填好，无需处理！")
        return

    print(f"🤖 准备处理 {len(rows_to_process)} 条数据...")
    
    # 3. 分批处理 (每批 30 条)
    batch_size = 30
    updates = {} # 暂存结果 {index: canonical_value}
    
    data_list = []
    for idx, row in rows_to_process.iterrows():
        data_list.append({
            "index": idx, # 记住原始行号
            "alias": row["Alias"],
            "advice": row["Advice"]
        })

    for i in tqdm(range(0, len(data_list), batch_size), desc="GPT Cleaning"):
        batch = data_list[i : i + batch_size]
        
        gpt_input = [{"alias": item["alias"], "advice": item["advice"]} for item in batch]
        
        # 调用 API
        gpt_res = ask_gpt_batch(gpt_input, api_key)
        
        # 解析结果并决定填什么
        for item in batch:
            alias = item["alias"]
            idx = item["index"]
            
            adv_id = df.at[idx, "Adviced_ID"]
            
            if alias in gpt_res:
                res = gpt_res[alias]
                
                if not res.get("is_company", False):
                    updates[idx] = "0"
                else:
                    if df.at[idx, "Advice"] and df.at[idx, "Adviced_ID"] and res.get("matches_advice", False):
                        updates[idx] = df.at[idx, "Adviced_ID"]
                    else:
                        updates[idx] = res.get("clean_name", alias)

    print("💾 正在保存结果...")
    for idx, val in updates.items():
        df.at[idx, "Canonical_Name"] = val
        
    df.to_csv(csv_path, index=False, encoding="utf-8-sig")
    cute_box(
        f"✅ 自动填写完成！已更新 {len(updates)} 行", 
        f"自動入力完了！{len(updates)} 行を更新しました", 
        "🎉"
    )

def step4():
    import pandas as _pd

    # 1) 读 CSV
    df = _pd.read_csv(BASE_DIR / "result.csv", dtype=str).fillna("")

    # 2) 准备输出行：注意这里给每一条都加上 value=1
    rows = []
    for _, r in tqdm(df.iterrows(), desc="生成邻接表", total=len(df)):
        comps = [r[f"company_{i}"] 
                 for i in range(1, MAX_COMP_COLS+1) 
                 if r[f"company_{i}"].strip()]
        for a, b in itertools.permutations(comps, 2):
            rows.append({
                "company_a": a,
                "company_b": b,
                "value": 1,
            })

    # 3) 构建完整 DataFrame
    out = _pd.DataFrame(rows)

    # 4) 写 adjacency list （只保留 a/b 两列）
    out[['company_a','company_b']].to_csv(
        BASE_DIR / "result_adjacency_list.csv",
        index=False, encoding="utf-8-sig"
    )
    cute_box(
        "Step4 已生成邻接表：result_adjacency_list.csv",
        "Step4 隣接リストを生成しました：result_adjacency_list.csv",
        "📋"
    )

    # ——— 生成带行列标题的 Pivot Table ———
    pivot = out.pivot_table(
        index="company_a",      # 行标签
        columns="company_b",    # 列标签
        values="value",         # 聚合字段
        aggfunc="sum",          # 把所有 value=1 累加
        fill_value=""           # 0 或 NaN 都显示空白
    )

    # 5) 导出带行/列标题的矩阵
    pivot.to_csv(
        BASE_DIR / "pivot_table.csv",
        encoding="utf-8-sig"
    )
    cute_box(
        "Step4 已生成透视表：pivot_table.csv",
        "Step4 ピボットテーブルを生成しました：pivot_table.csv",
        "📊"
    )
def main():
    # 1. 连接数据库
    mysql_url = ask_mysql_url()
    try:
        create_engine(mysql_url).connect().close()
        print("✅ 数据库连接成功 / データベース接続成功")
    except Exception as e:
        cute_box(f"数据库连接失败：{e}", f"データベース接続 失敗：{e}", "❌")
        sys.exit(1)
        
    # 2. 配置关键词模式
    configure_keywords()
    
    # 3. 主菜单循环
    while True:
        choice = choose()

        if choice == "1":
            # --- 阶段一：提取 ---
            step1()
            step2(mysql_url)
            
            # 标记：是否已经跑过 AI 清洗
            ai_cleaned_done = False

            # 跑完 Step 1-2 后，进入子菜单
            while True:
                print("\n" + "="*60)
                
                if not ai_cleaned_done:
                    # --- 状态 A：刚跑完提取，还没清洗 ---
                    print("🎉 [Step 1-2] 完成 / 完了")
                    print("   文件已生成: result_mapping_todo.csv")
                    print("   ファイル生成完了: result_mapping_todo.csv")
                    print("-" * 60)
                    print("👉 接下来建议做什么？/ 次のステップ：")
                    print("   [a] 🤖 运行 AI 自动名寄せ (强烈推荐) / AI自動名寄せを実行 [推奨]")
                    print("   [b] ⚠️ 跳过清洗，直接入库・分析・結果出力 / そのままDB登録へ進む・分析・結果出力")
                else:
                    # --- 状态 B：已经跑完 AI名寄せ ---
                    print("✨ [Step 2.5] AI名寄せ已完成 / AI名寄せ完了")
                    print("   请打开 result_mapping_todo.csv 简单检查一下，确认无误后继续。")
                    print("   名寄せ完了のresult_mapping_todo.csvを確認し、問題なければ次へ進んでください。")
                    print("-" * 60)
                    print("👉 下一步 / Next Step：")
                    print("   [b] 🚀 确认无误，执行入库・分析・結果出力 / 確認OK、DB登録・分析・結果出力")
                    print("   [a] 🔄 不满意，重跑 AI 清洗 / もう一度AIを実行")

                print("   [e] 👋 退出程序 / 一旦終了")
                print("="*60)
                
                sub_c = input("Input [a/b/e]: ").strip().lower()
                
                if sub_c == "a":
                    step_ai_autofill()
                    ai_cleaned_done = True # 标记为已清洗
                    
                elif sub_c == "b":
                    step3(mysql_url)
                    step4()
                    print("🎉 完成！ result_adjacency_list.csvやpivot_table.csvを確認してください〜")
                    sys.exit(0) # 全部完成，退出
                    
                elif sub_c == "e":
                    print("👋 Bye!")
                    sys.exit(0)

        elif choice == "2":
            # --- 阶段二：单独运行 AI 清洗 ---
            step_ai_autofill()
            print("\n✅ 完成。您可以选择 [3] 进行入库，或输入 [e] 退出。\n✅ 完成。 [3] でDB登録・分析・結果出力、もしくは [e] で終了。")
            # 这里可以不强制跳转，让用户自己选

        elif choice == "3":
            # --- 阶段三：入库与分析 ---
            step3(mysql_url)
            step4()
            print("🎉 所有任务已完成 / 全てのタスクが完了しました")
            sys.exit(0)

if __name__ == "__main__":
    main()

# 依赖 / Dependencies / 依存関係：python-docx, pandas, openpyxl, unicodedata, tqdm, spacy, fuzzywuzzy, python-Levenshtein
# 安装 / Install / インストール：pip install python-docx pandas openpyxl tqdm spacy fuzzywuzzy python-Levenshtein
# 模型下载 / Model Download / モデルダウンロード：python -m spacy download en_core_web_sm
# 原作者：杨天乐@关西大学 / Author: Shiame Yeung@Kansai University / 作成者：楊　天楽@関西大学
# 版本 / Version / バージョン：2025.06.05


import os
import re
import unicodedata
from docx import Document
import pandas as pd
from tqdm import tqdm

# ------------------ 用户可配置 ------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = SCRIPT_DIR  # 根目录就是脚本所在位置
OUTPUT_CSV = os.path.join(SCRIPT_DIR, 'keyword_hit.csv')  # 输出为 CSV，与 Step2/Step3 保持一致
PUB2COUNTRY_CSV = os.path.join(SCRIPT_DIR, 'Publisher to Country List.csv')
# -----------------------------------------------

# ------------------ 1. 读取 Publisher → Country 映射 ------------------
try:
    pub2country_df = pd.read_csv(PUB2COUNTRY_CSV, dtype=str, encoding='utf-8-sig')
    pub2country_map = {
        pub.strip(): country.strip()
        for pub, country in zip(pub2country_df['Publisher'], pub2country_df['Country'])
        if isinstance(pub, str)
    }
except Exception as e:
    print(f"❗ 无法读取 Publisher to Country List.csv: {e}")
    pub2country_map = {}
# -------------------------------------------------------------------

# ------------------ 2. 定义关键词根列表 ------------------
KEYWORD_ROOTS = [
    'partner','alliance','collaborat','cooper','cooperat','join','merger','acquisiti',
    'outsourc','invest','licens','integrat','coordinat','synergiz','associat',
    'confedera','federa','union','unit','amalgamat','conglomerat','combin',
    'buyout','companion','concur','concert','comply','complement','assist',
    'takeover','accession','procure','suppl','conjoint','support','adjust',
    'adjunct','patronag','subsid','affiliat','endors'
]
# --------------------------------------------------------

def clean_text(text: str) -> str:
    """
    去除 Unicode 控制字符（C 类别）并保留换行、制表符
    """
    return ''.join(c for c in text if unicodedata.category(c)[0] != 'C' or c in ('\n','\t'))

def _normalize(text: str) -> str:
    """
    归一化处理：
      1. 转小写
      2. 去掉两端空白
      3. 合并连续空白为单个空格
      4. 去掉首尾的破折号、冒号、引号
      5. 去掉常见标点符号（逗号、句号、分号、斜杠、括号等）
    """
    if not text:
        return ''
    t = text.lower().strip()
    t = re.sub(r'\s+', ' ', t)
    t = re.sub(r'^[\-:\"\']+', '', t)
    t = re.sub(r'[\-:\"\']+$', '', t)
    t = re.sub(r'[,.;/()]+', '', t)
    t = re.sub(r'\s+', ' ', t)
    return t.strip()

# ------------------ 3. 从索引区提取所有标题 ------------------
def extract_index_titles(paragraphs):
    """
    解析文档开头的索引，找到 Documents(n) 或 Document(n) 行，
    然后在其后连续抓取格式为 "数字. 空格 + 标题" 的行。
    对抓到的标题做归一化并去重，返回 [(num, title_raw, title_norm), ...] 列表。
    """
    paras_text = [p.text.strip() for p in paragraphs]
    titles_list = []

    # 匹配 "Document(n)" 或 "Documents(n)"
    doc_pattern = re.compile(r'Documents?\s*\(\s*(\d+)\s*\)', re.IGNORECASE)
    doc_count = None
    start_idx = None
    for idx, txt in enumerate(paras_text):
        m = doc_pattern.match(txt)
        if m:
            try:
                doc_count = int(m.group(1))
            except:
                pass
            start_idx = idx + 1
            break

    if doc_count is None or start_idx is None:
        return []

    # 从 start_idx 开始，匹配 "^(\d+)\.\s+(.*)$"
    pattern = re.compile(r'^(\d+)\.\s+(.*)$')
    seen_norm = set()
    for j in range(start_idx, len(paras_text)):
        txt = paras_text[j]
        if not txt:
            continue
        m2 = pattern.match(txt)
        if m2:
            num = int(m2.group(1))
            title_raw = m2.group(2).strip()
            title_norm = _normalize(title_raw)
            if title_norm in seen_norm:
                continue
            seen_norm.add(title_norm)
            titles_list.append((num, title_raw, title_norm))
            if len(titles_list) >= doc_count:
                break

    return sorted(titles_list, key=lambda x: x[0])

# ------------------ 4. 提取正文句子（包含 fallback 逻辑） ------------------
def extract_sentences_by_titles(filepath):
    """
    对单个 .docx 文件，先提取索引区标题，再逐篇定位正文：
      1. 如果能抓到 index_titles，则用归一化匹配索引标题到正文区，
         按顺序找到 Publisher、Body、Classification/Notes 区间并提取句子，
         保存 Title/Publisher/Country。
      2. 如果 index_titles 为空，则回退到旧逻辑：提取所有 Body→(Notes/Classification) 区对，
         句子记录的 Title/Publisher/Country 均置为空。
    返回一个 records 列表，每条为 dict：
      { 'Title','Publisher','Country','Sentence','Hit_Count','Matched_Keywords' }
    """
    records = []
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"❗ 无法读取文件 {filepath}: {e}")
        return []

    paras = doc.paragraphs
    paras_text = [p.text.strip() for p in paras]
    paras_norm = [_normalize(t) for t in paras_text]

    # 4.1 提取索引标题
    index_titles = extract_index_titles(paras)

    # 4.2 如果没有索引标题，则进入回退逻辑：提取所有 Body→Notes/Classification 对
    if not index_titles:
        # 找到所有 (body_start, body_end) 对
        pairs = []
        for idx, txt in enumerate(paras_text):
            if txt == 'Body':
                for k in range(idx + 1, len(paras_text)):
                    if paras_text[k] in ('Notes', 'Classification'):
                        pairs.append((idx + 1, k))
                        break
        # 如果没有任何 Body→Notes/Classification 区间，直接返回空
        if not pairs:
            return []
        for (bs, be) in pairs:
            full_text = []
            for k in range(bs, be):
                line = paras_text[k]
                if not line:
                    continue
                cleaned = clean_text(line)
                if cleaned:
                    full_text.append(cleaned)
            article_text = ' '.join(full_text)
            sentences = [s.strip() for s in re.split(r'\.\s*', article_text) if len(s.strip()) >= 5]
            for sent in sentences:
                hits = [root for root in KEYWORD_ROOTS if root in sent.lower()]
                records.append({
                    'Title': '',
                    'Publisher': '',
                    'Country': '',
                    'Sentence': sent,
                    'Hit_Count': len(hits),
                    'Matched_Keywords': '; '.join(hits)
                })
        return records

    # 4.3 如果有索引标题，则逐条匹配正文区
    for num, title_raw, title_norm in index_titles:
        match_idx = None
        for idx, norm_txt in enumerate(paras_norm):
            if norm_txt == title_norm:
                match_idx = idx
                break
        if match_idx is None:
            continue

        # 下一行是 Publisher
        pub_idx = match_idx + 1
        if pub_idx >= len(paras_text):
            continue
        publisher = paras_text[pub_idx].strip()
        country = pub2country_map.get(publisher, '')

        # 从 match_idx+2 开始找 Body
        body_start = None
        for k in range(match_idx + 2, len(paras_text)):
            if paras_text[k] == 'Body':
                body_start = k + 1
                break
        if body_start is None:
            continue

        # 从 body_start 开始找第一处 Notes 或 Classification
        body_end = None
        for k in range(body_start, len(paras_text)):
            if paras_text[k] in ('Classification', 'Notes'):
                body_end = k
                break
        if body_end is None:
            continue

        # 提取正文 [body_start, body_end)
        full_text = []
        for k in range(body_start, body_end):
            line = paras_text[k]
            if not line:
                continue
            cleaned = clean_text(line)
            if cleaned:
                full_text.append(cleaned)
        article_text = ' '.join(full_text)
        sentences = [s.strip() for s in re.split(r'\.\s*', article_text) if len(s.strip()) >= 5]
        for sent in sentences:
            hits = [root for root in KEYWORD_ROOTS if root in sent.lower()]
            records.append({
                'Title': title_raw,
                'Publisher': publisher,
                'Country': country,
                'Sentence': sent,
                'Hit_Count': len(hits),
                'Matched_Keywords': '; '.join(hits)
            })

    return records

# ------------------ 5. 主流程：遍历所有 docx 并汇总 ------------------
records_all = []
docx_files = []

for root, dirs, files in os.walk(ROOT_DIR):
    for fname in files:
        if not fname.endswith('.docx') or fname.startswith('~$'):
            continue
        path = os.path.join(root, fname)
        rel = os.path.relpath(path, ROOT_DIR).split(os.sep)
        if len(rel) >= 3:
            tier1, tier2, filename = rel[0], rel[1], rel[2]
        elif len(rel) == 2:
            tier1, tier2, filename = rel[0], '', rel[1]
        else:
            tier1, tier2, filename = '', '', rel[0]
        docx_files.append((path, tier1, tier2, filename))

for path, tier1, tier2, filename in tqdm(docx_files, desc="📄 处理 Word 文件", ncols=100):
    recs = extract_sentences_by_titles(path)
    for r in recs:
        r['Tier_1'] = tier1
        r['Tier_2'] = tier2
        r['Filename'] = filename
        records_all.append(r)

# 5.1 构建 DataFrame 并排序、输出 CSV
if records_all:
    df = pd.DataFrame(records_all, columns=[
        'Tier_1', 'Tier_2', 'Filename',
        'Title', 'Publisher', 'Country',
        'Sentence', 'Hit_Count', 'Matched_Keywords'
    ])
    # 原有排序：先按是否命中（Hit_Count>0），再按 Tier_1/Tier_2/Filename
    df['Hit_Flag'] = df['Hit_Count'].apply(lambda x: 1 if x > 0 else 0)
    df = df.sort_values(['Hit_Flag', 'Tier_1', 'Tier_2', 'Filename'],
                        ascending=[False, True, True, True]).reset_index(drop=True)
    df = df.drop(columns=['Hit_Flag'])

    df.to_csv(OUTPUT_CSV, index=False, encoding='utf-8-sig')
    print(f"✅ 全部处理完成，输出文件：{OUTPUT_CSV}")
else:
    print("⚠️ 没有找到任何有效的句子或 docx 文件为空。")
